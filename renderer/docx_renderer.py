
"""
renderer/docx_renderer.py
"""

import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

try:
    from renderer.screen_classifier import classify_screen
except Exception:
    class Dummy:
        def __init__(self,v): self.value=v
    def classify_screen(text):
        t=text.lower()
        if "talking head" in t: return Dummy("Talking Head")
        if "infographic" in t: return Dummy("Infographic")
        if "animation" in t or "line art" in t: return Dummy("Animation")
        if "activity" in t: return Dummy("Activity")
        if "reflection" in t: return Dummy("Reflection")
        if "question" in t: return Dummy("Question")
        if "image" in t: return Dummy("Image")
        return Dummy("Other")

SCREEN_COLORS={
"Talking Head":"D9EAD3",
"Infographic":"CFE2F3",
"Animation":"FFF2CC",
"Image":"EAD1DC",
"Activity":"F9CB9C",
"Reflection":"D0E0E3",
"Question":"F4CCCC",
"Other":"FFFFFF"
}

def shade_cell(cell,color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )

def style_header(cell,text):
    cell.text=""
    p=cell.paragraphs[0]
    p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    r=p.add_run(text)
    r.bold=True
    r.font.size=Pt(11)
    shade_cell(cell,"4F81BD")

def set_cell(cell,text,bold=False):
    cell.text=""
    r=cell.paragraphs[0].add_run(str(text))
    r.bold=bold
    r.font.size=Pt(10)

def classify_screen_non_presenter(text):
    t = text.lower()
    if "infographic" in t: return "Infographic"
    if "animation" in t or "line art" in t: return "Animation"
    if "activity" in t: return "Activity"
    if "reflection" in t: return "Reflection"
    if "question" in t: return "Question"
    if "image" in t: return "Image"
    return "Other"

def render_storyboard(input_json="output/generated/generated_storyboard.json",
                      output_docx="output/generated/generated_storyboard.docx"):
    input_json = Path(input_json)
    output_docx = Path(output_docx)
    with open(input_json,encoding="utf-8") as f:
        data=json.load(f)
    beats=data["beats"] if isinstance(data,dict) else data
    doc=Document()
    h=doc.add_heading("Generated Storyboard",1)
    h.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    table=doc.add_table(rows=1,cols=4)
    table.style="Table Grid"
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=table.rows[0].cells
    for c,t in zip(hdr,["Beat","Graphics / Animation","Text On Screen (OST)","Narrator Dialogue"]):
        style_header(c,t)
    for beat in beats:
        cells=table.add_row().cells
        set_cell(cells[0],beat.get("beat_id",""),True)
        visual=beat.get("visual","")
        
        # Read the planner's graphics_type decision directly from the beat
        graphics_type = beat.get("graphics_type")
        if graphics_type == "Talking Head":
            st = "Talking Head"
        elif graphics_type == "Other":
            # Classify but exclude Talking Head
            st = classify_screen_non_presenter(visual)
        elif graphics_type is not None:
            st = graphics_type
        else:
            # Backward compatibility fallback
            st = classify_screen(visual).value
            
        set_cell(cells[1],st+"\n\n"+visual)
        shade_cell(cells[1],SCREEN_COLORS.get(st,"FFFFFF"))
        set_cell(cells[2],beat.get("ost",""))
        set_cell(cells[3],beat.get("dialogue",""))
    output_docx.parent.mkdir(parents=True,exist_ok=True)
    doc.save(output_docx)

if __name__=="__main__":
    render_storyboard()
